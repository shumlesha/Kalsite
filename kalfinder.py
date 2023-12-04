from docx import Document
from fuzzywuzzy import fuzz
from fuzzywuzzy import process

#doc = Document("мудрый баобаб.docx")
#doc = Document("Collocopy.docx")
doc = Document("Thirdcollo.docx")

def read():
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip() != ""]


    questions = []
    answers = []
    current_answer = []

    for p in doc.paragraphs:

        if p.style.name.startswith('Heading'):
            if current_answer and "\n".join(current_answer).strip():

                answers.append("\n".join(current_answer))
                current_answer = []
            questions.append(p.text)
        else:

            current_answer.append(p.text)



    if current_answer:
        answers.append("\n".join(current_answer))

    '''print(len(questions), len(answers))
    for i, el in enumerate(questions):
        print(f"{i} {el} {answers[i]}")'''

    return dict(zip(questions, answers))

res = read()
q = list(res.keys())




def fastananlyzer(usermessage):

    return [[el[0], res[el[0]]] for el in process.extract(usermessage, q, limit=3)]
